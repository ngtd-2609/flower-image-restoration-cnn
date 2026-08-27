"""Finalize verified TOC numbers and replace one stale embedded EDA chart."""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "docs" / "REPORT_FINAL.docx"
CLASS_CHART = ROOT / "figures" / "eda" / "class_distribution.png"


def main() -> None:
    document = Document(REPORT)
    toc = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("CHƯƠNG 1. GIỚI THIỆU\t")
    )
    updates = {21: "39\n", 23: "40\n", 25: "41\n", 27: "43"}
    for run_index, value in updates.items():
        toc.runs[run_index].text = value

    abbreviation_table = document.tables[2]
    for row in abbreviation_table.rows:
        if row.cells[1].text.strip() == "Application Programming Interface":
            row.cells[0].text = "API"
            break

    replacements = {
        "Bộ test cục bộ thu thập 18 kiểm thử: 17 pass và 1 AppTest skip":
            "Bộ test cục bộ thu thập 22 kiểm thử: 21 pass và 1 AppTest skip",
        "17 test pass và 1 AppTest skip": "21 test pass và 1 AppTest skip",
        "File chất lượng ảnh hiện chỉ có 13 cột và không được đổi tên thành classification results.":
            "Bảng chất lượng ảnh kế thừa không thay thế 26.950 prediction CNN, 245 dòng per-class, "
            "66 dòng thống kê paired và manifest FULL_RUN.",
    }
    for paragraph in document.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if "classification_results_49_conditions.csv" in cell.text:
                    cell.text = cell.text.replace(
                        "classification_results_49_conditions.csv",
                        "results/final/condition_metrics.csv",
                    )
                if "results/error_analysis.csv" in cell.text:
                    cell.text = cell.text.replace(
                        "results/error_analysis.csv",
                        "results/final/error_analysis.csv",
                    )
    document.save(REPORT)

    # image2.png is the class-distribution figure, identified from document relationships.
    descriptor, temporary_name = tempfile.mkstemp(suffix=".docx", dir=REPORT.parent)
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        with ZipFile(REPORT, "r") as source, ZipFile(
            temporary, "w", compression=ZIP_DEFLATED, compresslevel=9
        ) as target:
            for info in source.infolist():
                data = (
                    CLASS_CHART.read_bytes()
                    if info.filename == "word/media/image2.png"
                    else source.read(info.filename)
                )
                target.writestr(info, data)
        temporary.replace(REPORT)
    finally:
        if temporary.exists():
            temporary.unlink()
    print(REPORT)


if __name__ == "__main__":
    main()
