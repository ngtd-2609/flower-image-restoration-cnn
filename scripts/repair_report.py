"""Apply targeted, layout-safe corrections to the existing final report."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "docs" / "REPORT_FINAL.docx"


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, value: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def size_table(table, widths: list[int], font_size: float = 8.5) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


def replace_in_paragraph(paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    full = paragraph.text.replace(old, new)
    paragraph.clear()
    paragraph.add_run(full)


def main() -> None:
    document = Document(REPORT)

    # Keep one evidence-status callout on the cover and remove the duplicate red line.
    status_table = document.tables[0]
    status_cell = status_table.cell(0, 0)
    status_cell.text = (
        "NOT_SUBMISSION_READY — Source, split, protocol, kiểm thử và Streamlit standalone đã được "
        "nâng cấp. Raw dataset, checkpoint, results/final, thông tin hành chính và deploy evidence "
        "chưa có; không công bố metric CNN hoặc tự nhận 95+."
    )
    status_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    status_paragraph = status_cell.paragraphs[0]
    status_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in status_paragraph.runs:
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(156, 64, 0)
    duplicate_status = document.paragraphs[8]
    duplicate_status._element.getparent().remove(duplicate_status._element)

    # Deterministic static TOC with the page map from the verified 44-page render.
    toc = next(paragraph for paragraph in document.paragraphs if paragraph.text.startswith("CHƯƠNG 1. GIỚI THIỆU\n"))
    toc.clear()
    toc.paragraph_format.line_spacing = 1.0
    toc.paragraph_format.space_before = Pt(0)
    toc.paragraph_format.space_after = Pt(0)
    toc.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.1), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    entries = [
        ("CHƯƠNG 1. GIỚI THIỆU", 8),
        ("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", 10),
        ("CHƯƠNG 3. DỮ LIỆU VÀ PHÂN TÍCH KHÁM PHÁ", 18),
        ("CHƯƠNG 4. PHƯƠNG PHÁP ĐỀ XUẤT", 23),
        ("CHƯƠNG 5. THIẾT KẾ THỰC NGHIỆM", 27),
        ("CHƯƠNG 6. KẾT QUẢ CHẤT LƯỢNG ẢNH", 29),
        ("CHƯƠNG 7. KẾT QUẢ NHẬN DIỆN MOBILENETV2", 33),
        ("CHƯƠNG 8. PHÂN TÍCH VÀ THẢO LUẬN", 34),
        ("CHƯƠNG 9. PHÂN TÍCH LỖI", 36),
        ("CHƯƠNG 10. ỨNG DỤNG VÀ TRIỂN KHAI", 37),
        ("CHƯƠNG 11. KẾT LUẬN", 40),
        ("CHƯƠNG 12. HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN", 41),
        ("TÀI LIỆU THAM KHẢO", 42),
        ("PHỤ LỤC A–D", 44),
    ]
    for index, (label, page) in enumerate(entries):
        run = toc.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.add_tab()
        page_run = toc.add_run(str(page))
        page_run.font.size = Pt(9)
        if index < len(entries) - 1:
            page_run.add_break(WD_BREAK.LINE)

    # Improve the three densest late-report tables without altering report structure.
    size_table(document.tables[17], [2100, 4300, 2960], font_size=8.5)
    size_table(document.tables[18], [1900, 3350, 4110], font_size=8.5)
    size_table(document.tables[19], [1750, 2450, 2450, 2710], font_size=8.5)

    for paragraph in document.paragraphs:
        replace_in_paragraph(
            paragraph,
            "Bộ test cục bộ đã qua 17 kiểm thử unit",
            "Bộ test cục bộ đã qua 21 kiểm thử; 1 AppTest được skip có lý do trong runtime thiếu Streamlit",
        )
        replace_in_paragraph(
            paragraph,
            "17 unit test",
            "21 kiểm thử PASS và 1 AppTest skip có lý do",
        )
        replace_in_paragraph(
            paragraph,
            "Bootstrap 95% resample theo ảnh; McNemar dùng bảng bất đồng trên cùng cặp prediction.",
            "Bootstrap 95% resample theo ảnh; exact McNemar dùng bảng bất đồng trên cùng cặp prediction "
            "và p-value được hiệu chỉnh Holm trên family 33 so sánh enhanced–degraded.",
        )
        replace_in_paragraph(
            paragraph,
            "nhiều so sánh nên hiệu chỉnh p-value.",
            "33 so sánh McNemar phải dùng Holm correction; báo đồng thời p-value raw và adjusted.",
        )

    document.core_properties.comments = (
        "Targeted evidence-status, TOC, table-layout and statistical-protocol corrections; "
        "no fabricated FULL_RUN results."
    )
    document.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
