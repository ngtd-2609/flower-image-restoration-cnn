from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS, FIG = ROOT / "docs", ROOT / "figures" / "report_final"
DOCS.mkdir(exist_ok=True); FIG.mkdir(parents=True, exist_ok=True)

facts = json.loads((ROOT / "artifacts/canonical_facts.json").read_text(encoding="utf-8"))
metrics = pd.read_csv(ROOT / "results/final/condition_metrics.csv")
per_class = pd.read_csv(ROOT / "results/final/per_class_metrics.csv")
stats = pd.read_csv(ROOT / "results/final/statistical_tests.csv")
pairs = pd.read_csv(ROOT / "results/final/top_confusion_pairs.csv")
tuning = pd.read_csv(ROOT / "results/validation_tuning_best.csv")
history = pd.read_csv(ROOT / "artifacts/training/history.csv")

clean = metrics.loc[metrics.image_type.eq("clean")].iloc[0]
degraded = metrics.loc[metrics.image_type.eq("degraded")]
enhanced = metrics.loc[metrics.image_type.eq("enhanced")]
best_enh = enhanced.loc[enhanced.macro_f1.idxmax()]
worst_deg = degraded.loc[degraded.macro_f1.idxmin()]


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "Cập nhật trường trong Microsoft Word"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end): run._r.append(node)


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = str(h); shade(c, "17365D")
        for r in c.paragraphs[0].runs: r.font.color.rgb = RGBColor(255,255,255); r.font.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths: cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return t


def picture(doc, path: Path, caption: str, width=6.25):
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        c = doc.add_paragraph(caption); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.style = "Caption"


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


# Final-result figures, generated only from results/final.
plot = metrics.copy()
plot["label"] = plot.condition_id.str.replace("__", " / ", regex=False)
top = plot.sort_values("macro_f1", ascending=False).head(12).sort_values("macro_f1")
fig, ax = plt.subplots(figsize=(9, 5.3)); ax.barh(top.label, top.macro_f1, color="#2A9D8F")
ax.axvline(clean.macro_f1, color="#E76F51", linestyle="--", label="Clean"); ax.set(xlabel="Macro F1", title="12 điều kiện có Macro F1 cao nhất")
ax.legend(); fig.tight_layout(); fig.savefig(FIG / "top_conditions.png", dpi=190); plt.close(fig)

agg = metrics.groupby(["image_type", "degradation"], as_index=False).macro_f1.mean()
fig, ax = plt.subplots(figsize=(8.5, 4.8))
for kind, frame in agg.groupby("image_type"):
    ax.plot(frame.degradation, frame.macro_f1, marker="o", label=kind)
ax.set(ylabel="Macro F1 trung bình", title="Nhận diện theo loại ảnh và cơ chế suy giảm"); ax.tick_params(axis="x", rotation=25); ax.legend()
fig.tight_layout(); fig.savefig(FIG / "macro_f1_by_type.png", dpi=190); plt.close(fig)

doc = Document()
section = doc.sections[0]; section.top_margin = Cm(2.0); section.bottom_margin = Cm(1.8); section.left_margin = Cm(2.3); section.right_margin = Cm(2.0)
styles = doc.styles
styles["Normal"].font.name = "Aptos"; styles["Normal"].font.size = Pt(10.5)
for name, size, color in (("Title", 24, "17365D"), ("Heading 1", 17, "17365D"), ("Heading 2", 13, "2A9D8F"), ("Heading 3", 11, "E76F51")):
    styles[name].font.name = "Aptos Display"; styles[name].font.size = Pt(size); styles[name].font.color.rgb = RGBColor.from_string(color)
    styles[name].font.bold = True

header = section.header.paragraphs[0]; header.text = "NHÓM 7  |  XỬ LÝ ẢNH  |  MOBILEnetV2"; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run("Trang "); add_field(footer, "PAGE")

# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("BÀI TẬP LỚN XỬ LÝ ẢNH").bold = True
doc.add_paragraph("\n")
p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("ĐÁNH GIÁ VÀ NÂNG CAO ĐỘ CHÍNH XÁC\nNHẬN DIỆN 5 LOẠI HOA\nTRONG ĐIỀU KIỆN ẢNH SUY GIẢM")
doc.add_paragraph("\n")
table(doc, ["Vai trò", "Thông tin"], [
    ["Giảng viên", facts["instructor"]], ["Nhóm", facts["group"]],
    ["Thành viên", "24100358 — Nguyễn Tùng Dương\n24100065 — Trịnh Ngọc Nga\n24106898 — Trương Việt Thành"],
    ["Trạng thái", "FULL_RUN_COMPLETE — DEPLOY_READY_BUT_NOT_DEPLOYED"], ["Ngày", "27/08/2026"],
])
doc.add_paragraph("Báo cáo này chỉ sử dụng số liệu từ pipeline CNN chạy thật; mọi surrogate và metric của Version A đã bị loại.").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

doc.add_heading("Tóm tắt", level=1)
doc.add_paragraph(
    f"Nghiên cứu đánh giá một MobileNetV2 cố định trên 3.670 ảnh thuộc năm lớp hoa. Dữ liệu được audit toàn phần và chia theo nhóm SHA-256 thành 2.571/549/550 ảnh, không giao đường dẫn hay hash. Một mô hình được huấn luyện hai giai đoạn; tham số xử lý ảnh chỉ được chọn trên Validation. Test được đánh giá trên đúng 49 điều kiện, tạo 26.950 dự đoán. Ảnh clean đạt Accuracy {clean.accuracy:.4f} và Macro F1 {clean.macro_f1:.4f}. Điều kiện enhanced tốt nhất là `{best_enh.condition_id}` với Macro F1 {best_enh.macro_f1:.4f}; điều kiện degraded yếu nhất là `{worst_deg.condition_id}` với Macro F1 {worst_deg.macro_f1:.4f}. Kết quả cho thấy enhancement không mặc nhiên cải thiện phân loại; hiệu quả phụ thuộc cơ chế suy giảm, mức độ và dấu hiệu lớp."
)
doc.add_heading("Đóng góp chính", level=2)
for text in (
    "Hợp nhất ba phiên bản ở cấp artifact, có inventory, scorecard và provenance cho từng quyết định.",
    "Audit dữ liệu bằng full decode, EXIF, mode/format gốc, kích thước và SHA-256; chống leakage theo nhóm trùng.",
    "Giao thức 49 điều kiện cố định với một CNN, cùng Test split và seed tái lập.",
    "Tuning chỉ trên Validation với thứ tự Macro F1 → SSIM → latency; cấu hình được khóa bằng checksum.",
    "Phân tích cặp, bootstrap CI, McNemar Holm, metric theo lớp và truy vết lỗi đến từng ảnh.",
): bullet(doc, text)

doc.add_heading("Mục lục", level=1); add_field(doc.add_paragraph(), "TOC \\o \"1-3\" \\h \\z \\u")
doc.add_page_break()

doc.add_heading("1. Bối cảnh và câu hỏi nghiên cứu", level=1)
doc.add_paragraph("Ảnh ngoài thực tế thường thiếu sáng, nhiễu, mất nét hoặc lệch màu. Một bộ phân loại tốt trên ảnh sạch có thể suy giảm mạnh khi phân phối đầu vào thay đổi. Dự án tách rõ hai vai trò: CNN thực hiện nhận diện; thuật toán xử lý ảnh là biến can thiệp trước CNN.")
for q in (
    "Suy giảm nào làm giảm Accuracy/Macro F1 mạnh nhất?", "Enhancement nào phục hồi nhận diện, và ở mức nào?",
    "PSNR/SSIM/Delta E có đồng biến với Macro F1 không?", "Lớp và cặp nhầm nào nhạy nhất?",
): bullet(doc, q)
doc.add_heading("1.1 Nguyên tắc chống thổi phồng kết quả", level=2)
doc.add_paragraph("Mọi kết luận dùng `results/final`. Không lấy số từ README, slide hoặc Word làm nguồn. Trạng thái triển khai công khai được giữ là chưa triển khai cho tới khi có URL và ảnh chụp xác minh độc lập.")

doc.add_heading("2. Hợp nhất Best-of-Three", level=1)
table(doc, ["Phiên bản", "Vai trò sau hợp nhất", "Quyết định"], [
    ["A", "Taxonomy lỗi và khung trình bày", "Loại surrogate/model/metrics"],
    ["B", "Nền kỹ thuật: src, tests, validator, Streamlit", "Chọn làm backbone"],
    ["C", "Bố cục báo cáo, slide, phân công", "Chọn cấu trúc nội dung; loại FastAPI legacy"],
])
doc.add_paragraph("Điểm baseline có chủ ý bảo thủ: A 44/100, B 73/100, C 65/100. Bảng đầy đủ nằm trong `artifacts/merge_audit/VERSION_SCORING.md`; provenance nằm trong `ARTIFACT_PROVENANCE.csv`.")

doc.add_heading("3. Dữ liệu và kiểm toán", level=1)
class_counts = facts["class_counts"]
table(doc, ["Lớp", "Số ảnh"], [[k, class_counts[k]] for k in facts["class_order"]])
doc.add_paragraph("Có 3 nhóm exact duplicate và 1 nhóm duplicate khác nhãn. Dự án không âm thầm xóa dữ liệu: các bản trùng được ghi nhận và buộc nằm cùng split bằng SHA-256. Không có ảnh lỗi giải mã trong bộ dữ liệu hiện được cung cấp.")
picture(doc, ROOT / "figures/eda/class_distribution.png", "Hình 1. Phân bố năm lớp sau audit")
picture(doc, ROOT / "figures/eda/sample_classes.png", "Hình 2. Ảnh đại diện theo lớp")

doc.add_heading("4. Chia tập và chống leakage", level=1)
table(doc, ["Split", "Số ảnh", "Tỷ lệ"], [[k.title(), v, f"{100*v/3670:.2f}%"] for k,v in facts["split_counts"].items()])
doc.add_paragraph("Grouped-stratified split dùng seed 42. Ba cặp giao đường dẫn và ba cặp giao SHA-256 đều bằng 0. Augmentation và degradation chỉ được áp dụng sau split. Validation dùng chọn mô hình/tham số; Test chỉ dùng đánh giá sau khóa.")

doc.add_heading("5. Tiền xử lý và mô hình", level=1)
doc.add_paragraph(f"Kiến trúc: {facts['model']}. Hợp đồng đầu vào: {facts['preprocessing']}.")
doc.add_heading("5.1 Huấn luyện hai giai đoạn", level=2)
doc.add_paragraph(f"Lịch sử có {len(history)} epoch thực tế: {sum(history.stage.eq('head'))} head và {sum(history.stage.eq('fine_tune'))} fine-tune. Checkpoint tốt nhất được lưu ở định dạng Keras. SHA-256: `{facts['model_checksum']}`.")
picture(doc, ROOT / "artifacts/training/learning_curves.png", "Hình 3. Loss và accuracy theo epoch")

doc.add_heading("6. Suy giảm và enhancement", level=1)
table(doc, ["Suy giảm", "Mức", "Enhancement"], [
    ["Low light", "light/medium/strong", "Gamma correction, CLAHE"],
    ["Gaussian noise", "light/medium/strong", "Gaussian, bilateral"],
    ["Salt & pepper", "light/medium/strong", "Median, Gaussian"],
    ["Gaussian blur", "light/medium/strong", "Unsharp mask, sharpening"],
    ["Color cast", "light/medium/strong", "RGB balance, HSV, LAB"],
])
picture(doc, ROOT / "figures/degradation/degradation_grid.png", "Hình 4. Minh họa năm cơ chế suy giảm")
picture(doc, ROOT / "figures/enhancement/enhancement_grid.png", "Hình 5. Ví dụ enhancement theo cơ chế")

doc.add_heading("7. Tuning trên Validation", level=1)
doc.add_paragraph("Mỗi ứng viên được đánh giá trên toàn bộ 549 ảnh Validation. Xếp hạng theo Macro F1 giảm dần; nếu hòa dùng SSIM giảm dần rồi latency tăng dần. Test path/hash không xuất hiện trong artifact tuning. Sau lựa chọn, tham số và model checksum được khóa trước Test.")
display_cols = [c for c in ["degradation", "level", "method", "macro_f1", "ssim", "latency_ms_per_image"] if c in tuning.columns]
table(doc, display_cols, [[f"{v:.4f}" if isinstance(v, float) else v for v in row] for row in tuning[display_cols].head(12).itertuples(index=False, name=None)])

doc.add_heading("8. Giao thức 49 điều kiện", level=1)
doc.add_paragraph("Ma trận gồm 1 clean + 15 degraded + 33 enhanced. Mỗi dòng dùng cùng 550 ảnh Test, cùng model SHA-256 và cùng preprocessing. Tổng cộng 26.950 bản ghi dự đoán; per-class table có 245 dòng; statistical table có 66 dòng.")

doc.add_heading("9. Kết quả chính", level=1)
table(doc, ["Chỉ tiêu", "Giá trị"], [
    ["Clean Accuracy", f"{clean.accuracy:.4f}"], ["Clean Macro F1", f"{clean.macro_f1:.4f}"],
    ["Enhanced tốt nhất", f"{best_enh.condition_id} — {best_enh.macro_f1:.4f}"],
    ["Degraded yếu nhất", f"{worst_deg.condition_id} — {worst_deg.macro_f1:.4f}"],
])
picture(doc, FIG / "top_conditions.png", "Hình 6. Các điều kiện có Macro F1 cao nhất")
picture(doc, FIG / "macro_f1_by_type.png", "Hình 7. Macro F1 theo loại ảnh")
top_rows = metrics.sort_values("macro_f1", ascending=False).head(15)
table(doc, ["Condition", "Type", "Accuracy", "Macro F1", "SSIM", "Latency ms"], [[
    r.condition_id, r.image_type, f"{r.accuracy:.4f}", f"{r.macro_f1:.4f}", f"{r.ssim:.4f}", f"{r.inference_time_ms_per_image_mean:.2f}"
] for r in top_rows.itertuples()])

doc.add_heading("10. Chất lượng ảnh và đánh đổi", level=1)
doc.add_paragraph("PSNR/SSIM đo độ trung thành so với ảnh clean; Delta E 2000 đặc biệt hữu ích cho lệch màu. Mọi metric ảnh dùng toàn bộ 550 ảnh và cùng lưới pixel xác định stride 4 (56×56 trên đầu vào 224×224), được ghi trong manifest để tái lập. Macro F1 là metric nhận diện chính. Một ảnh có SSIM cao hơn vẫn có thể làm mất dấu hiệu phân biệt lớp; ngược lại enhancement có thể tạo ảnh kém trung thành hơn nhưng giúp CNN.")
table(doc, ["Loại ảnh", "PSNR TB", "SSIM TB", "Macro F1 TB"], [[k, f"{g.psnr.mean():.3f}", f"{g.ssim.mean():.4f}", f"{g.macro_f1.mean():.4f}"] for k,g in metrics.groupby("image_type")])

doc.add_heading("11. Kiểm định thống kê", level=1)
doc.add_paragraph("Mỗi enhanced condition được so sánh cặp với degraded baseline trên cùng ảnh. Bootstrap báo khoảng tin cậy 95% cho chênh lệch Accuracy/Macro F1; McNemar kiểm tra thay đổi đúng–sai và p-value được hiệu chỉnh Holm trong họ 33 so sánh.")
sig = stats.loc[stats.mcnemar_reject_h0_0_05.astype(str).str.lower().eq("true")]
doc.add_paragraph(f"Có {len(sig)} hàng kiểm định bác bỏ H0 sau Holm trong tổng {len(stats)} hàng metric. Diễn giải tập trung cả độ lớn chênh lệch và CI, không chỉ p-value.")

doc.add_heading("12. Phân tích lỗi", level=1)
doc.add_paragraph("Error analysis được suy ra trực tiếp từ predictions đã lưu, tránh chạy suy luận lần hai. Mỗi hàng có đường dẫn, SHA-256, condition_id, nhãn thật, dự đoán clean/degraded/enhanced và confidence delta. Các nhóm gồm always_wrong, clean_correct_degraded_wrong, recovered_by_enhancement, harmed_by_enhancement, confidence_increased_still_wrong và enhanced_still_wrong.")
if not pairs.empty:
    cols = list(pairs.columns[:6]); table(doc, cols, [list(row) for row in pairs[cols].head(12).itertuples(index=False, name=None)])

doc.add_heading("13. Ứng dụng Streamlit và triển khai", level=1)
doc.add_paragraph("Ứng dụng độc lập hỗ trợ một ảnh và batch, chọn cơ chế/mức/enhancement, hiển thị ảnh trước–sau, top-k xác suất và thời gian. App chỉ báo READY khi model, metadata và checksum hợp lệ. Dockerfile và hướng dẫn deploy đã có; tuy nhiên chưa có URL Streamlit công khai nên trạng thái canonical là DEPLOY_READY_BUT_NOT_DEPLOYED.")

doc.add_heading("14. Hạn chế và đe dọa tính hợp lệ", level=1)
for text in (
    "Bộ dữ liệu chỉ có năm lớp hoa và có một nhóm duplicate khác nhãn; khả năng khái quát ngoài miền chưa được chứng minh.",
    "Degradation là mô phỏng có kiểm soát, không bao phủ toàn bộ camera/ánh sáng ngoài thực tế.",
    "Training/evaluation chạy CPU trong môi trường bàn giao; latency không đại diện GPU hoặc cloud.",
    "Tuning nhiều ứng viên có rủi ro thích nghi Validation; Test chỉ mở sau khóa giúp giảm nhưng không loại bỏ hoàn toàn rủi ro.",
    "Chưa có xác minh deployment công khai; không tuyên bố hoàn tất phần này.",
): bullet(doc, text)

doc.add_heading("15. Tái lập và kiểm định", level=1)
doc.add_paragraph("Lệnh chuẩn: `python scripts/run_full_pipeline.py --retrain`; kiểm định: `pytest -q`, `python scripts/check_notebook.py --require-full-run`, `python scripts/validate_project.py --require-full-run`. Gói nộp loại dữ liệu raw, cache và junction; SHA256SUMS cho phép xác minh integrity.")
table(doc, ["Artifact nguồn", "Vai trò"], [
    ["data/inventory.csv", "Raw audit"], ["splits/*.csv", "Split khóa"],
    ["models/model_metadata.json", "Model và checksum"], ["configs/locked_enhancement_params.json", "Tuning lock"],
    ["results/final/manifest.json", "Kết quả và checksum"], ["artifacts/canonical_facts.json", "Coherence cross-artifact"],
])

doc.add_heading("16. Kết luận", level=1)
doc.add_paragraph(f"Pipeline đã hoàn tất huấn luyện CNN thật và đánh giá full 49 điều kiện. Clean Macro F1 là {clean.macro_f1:.4f}; enhancement tốt nhất là {best_enh.condition_id} ({best_enh.macro_f1:.4f}). Kết luận quan trọng nhất là xử lý ảnh phải được chọn theo mục tiêu nhận diện và kiểm chứng trên dữ liệu tách biệt: cải thiện thị giác hoặc SSIM không bảo đảm cải thiện phân loại. Hệ thống hiện đủ bằng chứng local để tái lập và kiểm tra; deployment công khai còn là bước ngoài môi trường này.")

doc.add_heading("17. Tài liệu tham khảo", level=1)
for reference in (
    "[1] M. Sandler et al., “MobileNetV2: Inverted Residuals and Linear Bottlenecks,” CVPR, 2018, doi:10.1109/CVPR.2018.00474.",
    "[2] Z. Wang et al., “Image Quality Assessment: From Error Visibility to Structural Similarity,” IEEE TIP, 2004, doi:10.1109/TIP.2003.819861.",
    "[3] Q. McNemar, “Note on the Sampling Error of the Difference Between Correlated Proportions or Percentages,” Psychometrika, 1947, doi:10.1007/BF02295996.",
    "[4] B. Efron and R. Tibshirani, An Introduction to the Bootstrap, Chapman & Hall/CRC, 1993.",
    "[5] TensorFlow, Keras Applications — MobileNetV2 API, https://www.tensorflow.org/api_docs/python/tf/keras/applications/MobileNetV2.",
    "[6] Streamlit documentation, https://docs.streamlit.io/.",
): doc.add_paragraph(reference)

doc.add_heading("Phụ lục A — Truy vết Best-of-Three", level=1)
doc.add_paragraph("Inventory đầy đủ, component scorecard, decision matrix, remediation matrix và audit coherence được bàn giao trong `artifacts/`. Chúng ghi rõ nguồn A/B/C, quyết định chọn/loại, lý do, artifact đích và bằng chứng xác minh.")
doc.add_heading("Phụ lục B — Tuyên bố tính trung thực", level=1)
doc.add_paragraph("Không có số liệu placeholder trong báo cáo. Không dùng surrogate. Không tuyên bố URL deploy khi chưa có. Mọi số metric được đọc trực tiếp từ CSV final tại thời điểm tạo tài liệu.")

out = DOCS / "REPORT_FINAL.docx"
doc.save(out)
print(out)
